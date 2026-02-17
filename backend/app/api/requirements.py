from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Query
from fastapi.responses import JSONResponse
from io import BytesIO
from typing import Optional
import docx
from app.agents.requirement_reviewer import RequirementReviewer

router = APIRouter()
reviewer = RequirementReviewer()


@router.get("/categories")
async def get_categories():
    """获取可用品类列表"""
    try:
        categories = reviewer.get_available_categories()
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "data": categories
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e)
            }
        )


@router.get("/categories/{category_id}/fields")
async def get_category_fields(category_id: str, subtype_id: Optional[str] = None):
    """获取品类字段定义"""
    try:
        fields = reviewer.get_category_fields(category_id, subtype_id)
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "data": {
                    "category_id": category_id,
                    "subtype_id": subtype_id,
                    "fields": fields,
                    "total": len(fields)
                }
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e)
            }
        )


@router.post("/review-requirements")
async def review_requirements(
    file: UploadFile = File(...),
    category_id: Optional[str] = Form(None),
    subtype_id: Optional[str] = Form(None)
):
    """
    审查需求文档

    Args:
        file: 上传的文件（.docx或.txt）
        category_id: 品类ID（可选），如 'server', 'workstation'
        subtype_id: 子类型ID（可选），如 'gpu_ai_server'
    """
    try:
        # 读取文件内容
        content = await file.read()

        # 根据文件类型解析内容
        if file.filename.endswith('.docx'):
            text = _parse_docx(content)
        elif file.filename.endswith('.txt'):
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式，请上传.docx或.txt文件")

        # 使用审查智能体分析
        result = reviewer.review(text, category_id, subtype_id)

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "data": result,
                "filename": file.filename
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e)
            }
        )


@router.post("/review-requirements/text")
async def review_requirements_text(request: dict):
    """
    审查文本内容（无需上传文件）

    Request body:
    {
        "content": "需求文档内容...",
        "category_id": "server",  // 可选
        "subtype_id": "gpu_ai_server"  // 可选
    }
    """
    try:
        content = request.get("content", "")
        category_id = request.get("category_id")
        subtype_id = request.get("subtype_id")

        if not content or not content.strip():
            raise HTTPException(status_code=400, detail="内容不能为空")

        # 使用审查智能体分析
        result = reviewer.review(content, category_id, subtype_id)

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "data": result
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e)
            }
        )


def _parse_docx(content: bytes) -> str:
    """解析DOCX文件内容"""
    try:
        doc = docx.Document(BytesIO(content))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())

        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"解析DOCX文件失败: {str(e)}")
