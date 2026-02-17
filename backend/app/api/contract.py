from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from io import BytesIO
import docx
from app.agents.contract_analyzer import ContractAnalyzer

router = APIRouter()
analyzer = ContractAnalyzer()


@router.post("/contract-analysis")
async def analyze_contract(file: UploadFile = File(...)):
    """分析合同文档"""
    try:
        # 读取文件内容
        content = await file.read()

        # 根据文件类型解析内容
        if file.filename.endswith('.docx'):
            text = _parse_docx(content)
        elif file.filename.endswith('.txt'):
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式，请上传.docx或.txt文件")

        # 使用分析智能体分析
        result = analyzer.analyze(text)

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "data": result,
                "filename": file.filename
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e)
            }
        )


def _parse_docx(content: bytes) -> str:
    """解析DOCX文件内容"""
    try:
        doc = docx.Document(BytesIO(content))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())

        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"解析DOCX文件失败: {str(e)}")
